"""
agents/user_agent.py — UserAgent (Interaction + UI Layer)

Runs the Streamlit interface. Rich, fun, cutesy-academic aesthetic.
Fixes: retry history raw HTML bug, privacy false positives passed through from eval.
"""

import io
import time
import streamlit as st


class UserAgent:

    THEME_CSS = """
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Playfair+Display:ital,wght@0,600;0,700;1,500&family=DM+Sans:wght@300;400;500;600&family=DM+Mono:wght@400;500&display=swap');

    :root {
        --cream:      #fdf6ec;
        --warm-white: #fff9f2;
        --peach:      #f4a06e;
        --coral:      #e8623c;
        --deep-coral: #c94a28;
        --goldenrod:  #d4922a;
        --gold-light: #f0c96a;
        --sage:       #5a7a5e;
        --muted-sage: #8aac8e;
        --lavender:   #b8a9c9;
        --ink:        #2b2118;
        --ink-light:  #5c4a38;
        --card-bg:    rgba(255,249,242,0.95);
        --border:     rgba(212,146,42,0.22);
        --radius:     14px;
        --radius-lg:  20px;
    }

    html, body, [class*="css"] {
        font-family: 'DM Sans', sans-serif;
        color: var(--ink);
        background-color: var(--cream);
    }

    .main {
        background-image: radial-gradient(circle, rgba(212,146,42,0.12) 1px, transparent 1px);
        background-size: 28px 28px;
    }

    .main .block-container {
        max-width: 1440px;
        padding: 0 2rem 3rem;
    }

    ::-webkit-scrollbar { width: 6px; height: 6px; }
    ::-webkit-scrollbar-track { background: transparent; }
    ::-webkit-scrollbar-thumb { background: var(--peach); border-radius: 999px; }

    #MainMenu, footer, .stDeployButton { visibility: hidden; }

    .octo-header {
        display: flex;
        align-items: center;
        justify-content: space-between;
        padding: 1.8rem 0 1rem;
        border-bottom: 2px solid var(--border);
        margin-bottom: 0.5rem;
    }
    .octo-left { display: flex; align-items: center; gap: 1rem; }
    .octo-logo {
        font-size: 3.2rem;
        animation: float 3s ease-in-out infinite;
        display: inline-block;
    }
    @keyframes float {
        0%,100% { transform: translateY(0); }
        50%      { transform: translateY(-6px); }
    }
    .octo-title {
        font-family: 'Playfair Display', serif;
        font-size: 2.6rem;
        font-weight: 700;
        color: var(--coral);
        line-height: 1;
        margin: 0;
    }
    .octo-tagline {
        font-size: 0.78rem;
        color: var(--ink-light);
        letter-spacing: 0.1em;
        text-transform: uppercase;
        margin-top: 0.3rem;
        font-weight: 500;
    }
    .octo-badge-row { display: flex; gap: 0.5rem; flex-wrap: wrap; }
    .octo-badge {
        background: rgba(212,146,42,0.12);
        border: 1px solid var(--border);
        border-radius: 999px;
        padding: 0.2rem 0.75rem;
        font-size: 0.72rem;
        color: var(--goldenrod);
        font-weight: 600;
        letter-spacing: 0.05em;
        text-transform: uppercase;
    }
    .octo-badge.green { background: rgba(90,122,94,0.1); color: var(--sage); border-color: rgba(90,122,94,0.3); }
    .octo-badge.coral { background: rgba(232,98,60,0.1); color: var(--coral); border-color: rgba(232,98,60,0.3); }

    .panel-card {
        background: var(--card-bg);
        border: 1.5px solid var(--border);
        border-radius: var(--radius);
        padding: 1.3rem 1.5rem;
        margin-bottom: 1rem;
        box-shadow: 0 2px 14px rgba(43,33,24,0.06);
        position: relative;
        overflow: hidden;
    }
    .panel-card::before {
        content: '';
        position: absolute;
        top: 0; left: 0; right: 0;
        height: 3px;
        background: linear-gradient(90deg, var(--peach), var(--coral), var(--goldenrod));
    }
    .panel-label {
        font-family: 'Playfair Display', serif;
        font-size: 1rem;
        font-weight: 600;
        color: var(--goldenrod);
        letter-spacing: 0.02em;
        margin-bottom: 0.75rem;
        display: flex;
        align-items: center;
        gap: 0.4rem;
    }

    [data-testid="stFileUploader"] {
        border: 2px dashed var(--peach) !important;
        border-radius: var(--radius) !important;
        background: rgba(244,160,110,0.04) !important;
    }

    .conf-card {
        background: var(--card-bg);
        border: 1.5px solid var(--border);
        border-radius: var(--radius);
        padding: 1.3rem 1.5rem;
        margin-bottom: 1rem;
        position: relative;
        overflow: hidden;
    }
    .conf-card::before {
        content: '';
        position: absolute;
        top: 0; left: 0; right: 0;
        height: 3px;
        background: linear-gradient(90deg, var(--peach), var(--coral), var(--goldenrod));
    }
    .conf-number {
        font-family: 'Playfair Display', serif;
        font-size: 3rem;
        font-weight: 700;
        line-height: 1;
        margin-bottom: 0.2rem;
    }
    .conf-bar-track {
        background: rgba(212,146,42,0.15);
        border-radius: 999px;
        height: 10px;
        overflow: hidden;
        margin: 0.6rem 0 0.5rem;
    }
    .conf-bar-fill {
        height: 100%;
        border-radius: 999px;
        background: linear-gradient(90deg, var(--peach), var(--coral));
    }
    .quality-badge {
        display: inline-block;
        padding: 0.22rem 0.85rem;
        border-radius: 999px;
        font-size: 0.72rem;
        font-weight: 600;
        letter-spacing: 0.1em;
        text-transform: uppercase;
    }
    .badge-excellent { background: #d4edda; color: #1e5631; }
    .badge-good      { background: #dff0d8; color: #2d5a27; }
    .badge-fair      { background: #fef3cd; color: #7d5a00; }
    .badge-poor      { background: #fde0d8; color: #7d1f00; }
    .badge-empty     { background: #e9ecef; color: #495057; }

    .stat-grid {
        display: grid;
        grid-template-columns: repeat(auto-fill, minmax(110px, 1fr));
        gap: 0.55rem;
        margin: 0.5rem 0 0.8rem;
    }
    .stat-tile {
        background: rgba(212,146,42,0.07);
        border: 1px solid var(--border);
        border-radius: 10px;
        padding: 0.55rem 0.75rem;
        text-align: center;
    }
    .stat-tile-val {
        font-family: 'DM Mono', monospace;
        font-size: 1.05rem;
        font-weight: 500;
        color: var(--coral);
        display: block;
    }
    .stat-tile-lbl {
        font-size: 0.67rem;
        color: var(--ink-light);
        text-transform: uppercase;
        letter-spacing: 0.07em;
        margin-top: 0.1rem;
    }

    .loop-strip {
        display: flex;
        align-items: center;
        gap: 0;
        margin: 0.8rem 0 1rem;
        overflow-x: auto;
        padding: 0.2rem 0;
    }
    .loop-node {
        display: flex;
        flex-direction: column;
        align-items: center;
        gap: 0.2rem;
        flex-shrink: 0;
    }
    .loop-icon {
        width: 38px; height: 38px;
        border-radius: 50%;
        display: flex;
        align-items: center;
        justify-content: center;
        font-size: 1.1rem;
        border: 2px solid transparent;
    }
    .loop-icon.done  { background: var(--sage); border-color: #3d5e41; }
    .loop-icon.idle  { background: rgba(212,146,42,0.12); border-color: var(--border); }
    .loop-arrow {
        font-size: 0.9rem;
        color: var(--peach);
        padding: 0 0.15rem;
        flex-shrink: 0;
        margin-bottom: 1.2rem;
    }
    .loop-lbl {
        font-size: 0.6rem;
        color: var(--ink-light);
        text-transform: uppercase;
        letter-spacing: 0.05em;
        font-weight: 600;
    }

    .privacy-banner {
        background: linear-gradient(135deg, #fff5f2, #fff0eb);
        border: 2px solid var(--coral);
        border-radius: var(--radius);
        padding: 1rem 1.3rem;
        margin: 0.8rem 0;
        position: relative;
        overflow: hidden;
    }
    .privacy-banner::after {
        content: '🔐';
        position: absolute;
        right: 1rem; top: 50%;
        transform: translateY(-50%);
        font-size: 3.5rem;
        opacity: 0.07;
    }
    .privacy-banner h4 {
        color: var(--coral);
        margin: 0 0 0.3rem;
        font-family: 'Playfair Display', serif;
        font-size: 1rem;
    }
    .privacy-item {
        display: flex;
        justify-content: space-between;
        align-items: center;
        padding: 0.28rem 0;
        border-bottom: 1px dashed rgba(232,98,60,0.2);
        font-size: 0.83rem;
    }
    .privacy-item:last-child { border-bottom: none; }
    .priv-tag {
        background: rgba(232,98,60,0.1);
        border-radius: 6px;
        padding: 0.1rem 0.5rem;
        font-size: 0.72rem;
        font-family: 'DM Mono', monospace;
        color: var(--deep-coral);
    }

    .step-log {
        background: #1e160f;
        border-radius: var(--radius);
        padding: 1rem 1.2rem;
        font-family: 'DM Mono', monospace;
        font-size: 0.76rem;
        color: #d4b896;
        max-height: 240px;
        overflow-y: auto;
        line-height: 1.75;
        border: 1px solid rgba(212,146,42,0.15);
    }
    .step-log .log-line  { margin: 0; white-space: pre-wrap; word-break: break-word; }
    .step-log .log-warn  { color: #f4a06e; }
    .step-log .log-ok    { color: #8aac8e; }
    .step-log .log-sep   { color: #4a3828; }
    .step-log .log-agent { color: #d4922a; }

    .fun-quote {
        background: linear-gradient(135deg, rgba(244,160,110,0.12), rgba(212,146,42,0.08));
        border-left: 3px solid var(--peach);
        border-radius: 0 10px 10px 0;
        padding: 0.6rem 1rem;
        font-style: italic;
        font-size: 0.82rem;
        color: var(--ink-light);
        margin: 0.7rem 0;
    }
    .fun-quote strong { color: var(--coral); font-style: normal; }

    .strategy-chip {
        display: inline-flex;
        align-items: center;
        gap: 0.35rem;
        background: rgba(90,122,94,0.1);
        border: 1px solid rgba(90,122,94,0.3);
        border-radius: 999px;
        padding: 0.25rem 0.8rem;
        font-size: 0.78rem;
        color: var(--sage);
        font-weight: 500;
        font-family: 'DM Mono', monospace;
    }

    .stButton > button {
        background: linear-gradient(135deg, var(--coral), var(--deep-coral)) !important;
        color: white !important;
        border: none !important;
        border-radius: 10px !important;
        font-family: 'DM Sans', sans-serif !important;
        font-weight: 600 !important;
        padding: 0.55rem 1.4rem !important;
        letter-spacing: 0.02em !important;
    }
    .stButton > button:hover {
        background: linear-gradient(135deg, var(--goldenrod), var(--coral)) !important;
        transform: translateY(-1px) !important;
        box-shadow: 0 4px 14px rgba(232,98,60,0.3) !important;
    }

    .stTextArea textarea {
        font-family: 'DM Mono', monospace !important;
        font-size: 0.86rem !important;
        background: var(--warm-white) !important;
        border: 1.5px solid var(--border) !important;
        border-radius: 10px !important;
        color: var(--ink) !important;
        line-height: 1.65 !important;
    }

    [data-testid="stDownloadButton"] > button {
        background: transparent !important;
        color: var(--goldenrod) !important;
        border: 1.5px solid var(--border) !important;
        border-radius: 10px !important;
        font-weight: 500 !important;
    }
    [data-testid="stDownloadButton"] > button:hover {
        background: rgba(212,146,42,0.08) !important;
        border-color: var(--goldenrod) !important;
    }

    [data-testid="stExpander"] {
        border: 1.5px solid var(--border) !important;
        border-radius: var(--radius) !important;
        background: var(--card-bg) !important;
    }

    .empty-state { text-align: center; padding: 4rem 2rem; color: var(--ink-light); }
    .empty-state .big-emoji { font-size: 4rem; margin-bottom: 1rem; }
    .empty-state h3 {
        font-family: 'Playfair Display', serif;
        color: var(--goldenrod);
        margin: 0 0 0.5rem;
        font-size: 1.3rem;
    }
    .empty-state p { font-size: 0.85rem; margin: 0; line-height: 1.6; }

    .wc-bar-wrap { margin: 0.4rem 0 0.6rem; }
    .wc-label { font-size: 0.75rem; color: var(--ink-light); margin-bottom: 0.2rem; }
    .wc-bar-bg { background: rgba(90,122,94,0.15); border-radius: 999px; height: 6px; }
    .wc-bar-fill { height: 100%; border-radius: 999px; background: linear-gradient(90deg, var(--muted-sage), var(--sage)); }

    .sec-divider {
        display: flex; align-items: center; gap: 0.6rem;
        margin: 1.2rem 0 0.7rem;
        font-size: 0.72rem; color: var(--ink-light);
        text-transform: uppercase; letter-spacing: 0.1em; font-weight: 600;
    }
    .sec-divider::before, .sec-divider::after {
        content: ''; flex: 1; height: 1px; background: var(--border);
    }

    [data-testid="stImage"] img {
        border-radius: var(--radius) !important;
        border: 2px solid var(--border) !important;
        box-shadow: 0 4px 20px rgba(43,33,24,0.1) !important;
    }
    </style>
    """

    LOADING_PHRASES = [
        "🧠 DocAgent is thinking very hard...",
        "🔍 ImageAgent squinting at your pixels...",
        "⚡ StrategyAgent consulting the oracle...",
        "📝 Tesseract doing its best impression of a human...",
        "⚖️ EvalAgent judging every character ruthlessly...",
        "🔒 Privacy scanner on patrol...",
        "✨ Polishing the final output...",
    ]

    RESULT_QUOTES = {
        "EXCELLENT": ("🎉", "Absolutely nailed it.", "The agents are feeling smug right now."),
        "GOOD":      ("✅", "Pretty solid extraction.", "The retry loop earned its keep."),
        "FAIR":      ("🤔", "Good enough, but the image could be cleaner.", "Garbage in, less garbage out."),
        "POOR":      ("😬", "OCR tried its best. Image quality was rough.", "Consider a cleaner scan next time."),
    }

    def run(self):
        st.markdown(self.THEME_CSS, unsafe_allow_html=True)
        self._render_header()
        left, right = st.columns([1, 1.4], gap="large")
        with left:
            self._render_upload_panel()
        with right:
            self._render_output_panel()

    def _render_header(self):
        st.markdown("""
        <div class="octo-header">
          <div class="octo-left">
            <div class="octo-logo">🐙</div>
            <div>
              <div class="octo-title">OctoRead</div>
              <div class="octo-tagline">Agentic OCR · Intelligent Document Extraction</div>
            </div>
          </div>
          <div class="octo-badge-row">
            <span class="octo-badge">5-Agent Loop</span>
            <span class="octo-badge green">Adaptive Retry</span>
            <span class="octo-badge coral">Privacy Aware</span>
            <span class="octo-badge">Human-in-Loop</span>
          </div>
        </div>
        """, unsafe_allow_html=True)

    def _render_upload_panel(self):
        st.markdown('<div class="panel-label">📎 Upload Image</div>', unsafe_allow_html=True)
        uploaded = st.file_uploader(
            "Drop an image here",
            type=["png", "jpg", "jpeg", "bmp", "tiff", "webp"],
            label_visibility="collapsed",
        )

        if uploaded:
            st.image(uploaded, use_container_width=True)
            image_bytes = uploaded.read()
            size_kb = len(image_bytes) / 1024
            st.markdown(f"""
            <div style="display:flex;gap:0.5rem;margin:0.4rem 0 0.8rem;flex-wrap:wrap">
              <span class="strategy-chip">📁 {uploaded.name}</span>
              <span class="strategy-chip">💾 {size_kb:.1f} KB</span>
              <span class="strategy-chip">🎨 {uploaded.type.split("/")[1].upper()}</span>
            </div>
            """, unsafe_allow_html=True)

            st.markdown('<div class="sec-divider">options</div>', unsafe_allow_html=True)
            use_llm = st.toggle("✨ LLM Enhancement (Claude API)", value=False,
                                help="Post-processes OCR text with Claude. Requires API key.")
            if use_llm:
                api_key = st.text_input("Anthropic API Key", type="password", placeholder="sk-ant-...")
                if api_key:
                    st.session_state["anthropic_api_key"] = api_key

            st.markdown('<div class="sec-divider">agent loop</div>', unsafe_allow_html=True)
            self._render_loop_diagram("done" if "ocr_result" in st.session_state else "idle")
            st.markdown("")
            if st.button("🚀  Run Agentic OCR", use_container_width=True):
                self._run_agents(image_bytes, use_llm=use_llm)
        else:
            st.markdown("""
            <div class="empty-state">
              <div class="big-emoji">🖼️</div>
              <h3>Nothing here yet!</h3>
              <p>Upload a PNG, JPG, TIFF, or BMP of<br>
              typed text, a receipt, a document,<br>
              or anything with words on it.</p>
            </div>
            """, unsafe_allow_html=True)
            st.markdown('<div class="sec-divider">how it works</div>', unsafe_allow_html=True)
            self._render_loop_diagram("idle")

    def _render_loop_diagram(self, state: str):
        nodes = [("👁","Observe"),("⚡","Decide"),("🔧","Act"),("⚖️","Evaluate"),("🔁","Improve")]
        html = '<div class="loop-strip">'
        for i, (icon, lbl) in enumerate(nodes):
            cls = "done" if state == "done" else "idle"
            html += f'<div class="loop-node"><div class="loop-icon {cls}">{icon}</div><div class="loop-lbl">{lbl}</div></div>'
            if i < len(nodes) - 1:
                html += '<div class="loop-arrow">→</div>'
        html += '</div>'
        st.markdown(html, unsafe_allow_html=True)

    def _render_output_panel(self):
        if "ocr_result" not in st.session_state:
            st.markdown("""
            <div class="empty-state">
              <div class="big-emoji">📄</div>
              <h3>Awaiting Extraction</h3>
              <p>Upload an image and click<br><em>Run Agentic OCR</em> to begin.</p>
            </div>
            """, unsafe_allow_html=True)
            return

        r = st.session_state["ocr_result"]
        self._render_confidence_card(r)
        self._render_image_stats(r)
        self._render_result_quote(r.get("quality_label", "FAIR"))
        self._render_privacy_warnings(r.get("privacy_warnings", []))
        st.markdown('<div class="sec-divider">extracted text</div>', unsafe_allow_html=True)
        self._render_text_editor(r)
        st.markdown('<div class="sec-divider">agent internals</div>', unsafe_allow_html=True)
        self._render_decision_log(r.get("decision_log", []))
        self._render_attempt_history(r.get("all_attempts", []))
        st.markdown('<div class="sec-divider">export</div>', unsafe_allow_html=True)
        self._render_export(r)

    def _render_confidence_card(self, r: dict):
        conf = r.get("confidence", 0)
        quality = r.get("quality_label", "POOR") or "POOR"
        colors = {"EXCELLENT": "#1e5631", "GOOD": "#2d5a27", "FAIR": "#7d5a00", "POOR": "#8b2500"}
        fill_color = colors.get(quality, "#e8623c")
        badge_cls = f"badge-{quality.lower()}"
        strategy = r.get("strategy_used", "—")
        attempts = r.get("attempts", 1)
        reason = r.get("strategy_reason", "")
        st.markdown(f"""
        <div class="conf-card">
          <div class="panel-label">📊 OCR Confidence Score</div>
          <div style="display:flex;align-items:flex-end;gap:1rem;flex-wrap:wrap">
            <div class="conf-number" style="color:{fill_color}">{conf:.1f}%</div>
            <div style="padding-bottom:0.5rem"><span class="quality-badge {badge_cls}">{quality}</span></div>
          </div>
          <div class="conf-bar-track"><div class="conf-bar-fill" style="width:{conf}%"></div></div>
          <div style="display:flex;gap:0.5rem;align-items:center;flex-wrap:wrap;margin-top:0.3rem">
            <span class="strategy-chip">🧪 {strategy}</span>
            <span style="font-size:0.78rem;color:var(--ink-light)">{attempts} attempt{'s' if attempts != 1 else ''} · {r.get('word_count', 0)} words</span>
          </div>
          {f'<div style="font-size:0.75rem;color:var(--ink-light);margin-top:0.4rem;font-style:italic">"{reason}"</div>' if reason else ''}
        </div>
        """, unsafe_allow_html=True)

    def _render_image_stats(self, r: dict):
        f = r.get("image_features", {})
        blur = f.get("blur_score", 0)
        blur_d = f"{blur:.0f}" if isinstance(blur, float) else str(blur)
        contrast = f.get("contrast_score", 0)
        contrast_d = f"{contrast:.0f}" if isinstance(contrast, float) else str(contrast)
        w, h = f.get("width", 0), f.get("height", 0)
        st.markdown(f"""
        <div class="stat-grid">
          <div class="stat-tile"><span class="stat-tile-val">{blur_d}</span><div class="stat-tile-lbl">Sharpness</div></div>
          <div class="stat-tile"><span class="stat-tile-val">{contrast_d}</span><div class="stat-tile-lbl">Contrast</div></div>
          <div class="stat-tile"><span class="stat-tile-val">{f.get('noise_level','—')}</span><div class="stat-tile-lbl">Noise</div></div>
          <div class="stat-tile"><span class="stat-tile-val">{f.get('brightness','—')}</span><div class="stat-tile-lbl">Brightness</div></div>
          <div class="stat-tile"><span class="stat-tile-val">{f.get('document_type','—')}</span><div class="stat-tile-lbl">Doc Type</div></div>
          <div class="stat-tile"><span class="stat-tile-val" style="font-size:0.85rem">{w}×{h}</span><div class="stat-tile-lbl">Dimensions</div></div>
        </div>
        """, unsafe_allow_html=True)

    def _render_result_quote(self, quality: str):
        if quality not in self.RESULT_QUOTES:
            return
        emoji, headline, sub = self.RESULT_QUOTES[quality]
        st.markdown(f'<div class="fun-quote">{emoji} <strong>{headline}</strong> {sub}</div>', unsafe_allow_html=True)

    def _render_privacy_warnings(self, warnings: list):
        if not warnings:
            return
        items = "".join(f"""
        <div class="privacy-item">
          <span>⚠️ {w['type']}</span>
          <span class="priv-tag">{w['preview']} ×{w['count']}</span>
        </div>""" for w in warnings)
        st.markdown(f"""
        <div class="privacy-banner">
          <h4>🔐 Sensitive Data Detected</h4>
          <p style="font-size:0.8rem;color:#5c4a38;margin:0 0 0.5rem">
            Review carefully before sharing or downloading.
          </p>{items}
        </div>
        """, unsafe_allow_html=True)

    def _render_text_editor(self, r: dict):
        initial = st.session_state.get("edited_text", r["text"])
        edited = st.text_area("Edit before downloading:", value=initial, height=300,
                              label_visibility="visible", key="text_editor_area")
        st.session_state["edited_text"] = edited
        wc = len(edited.split()) if edited.strip() else 0
        bar_pct = min(wc / 5, 100)
        st.markdown(f"""
        <div class="wc-bar-wrap">
          <div class="wc-label">{wc} words in output</div>
          <div class="wc-bar-bg"><div class="wc-bar-fill" style="width:{bar_pct}%"></div></div>
        </div>""", unsafe_allow_html=True)

    def _render_decision_log(self, log: list):
        with st.expander("🤖 Agent Decision Log", expanded=False):
            if not log:
                st.caption("No log entries.")
                return
            lines_html = ""
            for line in log:
                l = line.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")
                if any(k in line for k in ("⚠️","Privacy","sensitive","Sensitive")):
                    cls = "log-warn"
                elif any(k in line for k in ("acceptable","complete","Stopping","best","✅")):
                    cls = "log-ok"
                elif "---" in line:
                    cls = "log-sep"
                elif any(k in line for k in ("DocAgent","ImageAgent","StrategyAgent","EvalAgent")):
                    cls = "log-agent"
                else:
                    cls = ""
                lines_html += f'<p class="log-line {cls}">{l}</p>\n'
            st.markdown(f'<div class="step-log">{lines_html}</div>', unsafe_allow_html=True)

    def _render_attempt_history(self, attempts: list):
        if len(attempts) <= 1:
            return
        with st.expander(f"🔁 Retry History — {len(attempts)} attempts", expanded=False):
            import pandas as pd
            rows = []
            for a in attempts:
                quality = "✅ GOOD" if a.confidence >= 65 else "⚠️ FAIR" if a.confidence >= 45 else "❌ POOR"
                rows.append({"#": a.attempt_num, "Strategy": a.strategy_name,
                             "Confidence": f"{a.confidence:.1f}%", "Words": a.word_count, "Quality": quality})
            df = pd.DataFrame(rows).set_index("#")
            st.dataframe(df, use_container_width=True)

    def _render_export(self, r: dict):
        text = st.session_state.get("edited_text", r["text"])
        c1, c2, c3 = st.columns(3)
        with c1:
            st.download_button("📄 Download .txt", data=text.encode("utf-8"),
                               file_name="octoread_output.txt", mime="text/plain",
                               use_container_width=True)
        with c2:
            try:
                st.download_button("📝 Download .docx", data=self._make_docx(text, r),
                                   file_name="octoread_output.docx",
                                   mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                   use_container_width=True)
            except Exception:
                st.caption("Install python-docx for Word export.")
        with c3:
            md = f"# OctoRead OCR Output\n\n**Confidence:** {r.get('confidence',0):.1f}%\n**Strategy:** {r.get('strategy_used','—')}\n**Words:** {r.get('word_count',0)}\n\n---\n\n{text}"
            st.download_button("📋 Download .md", data=md.encode("utf-8"),
                               file_name="octoread_output.md", mime="text/markdown",
                               use_container_width=True)

    def _run_agents(self, image_bytes: bytes, use_llm: bool = False):
        from agents.doc_agent import DocAgent
        placeholder = st.empty()
        with placeholder.container():
            pbar = st.progress(0, text=self.LOADING_PHRASES[0])
            for phrase, pct in zip(self.LOADING_PHRASES, [0.08, 0.22, 0.42, 0.60, 0.78, 0.91, 1.0]):
                pbar.progress(pct, text=phrase)
                time.sleep(0.28)
        try:
            agent = DocAgent()
            result = agent.run(image_bytes)
            conf = result.get("confidence", 0)
            result["quality_label"] = ("EXCELLENT" if conf >= 80 else "GOOD" if conf >= 65
                                       else "FAIR" if conf >= 45 else "POOR")
            if use_llm and st.session_state.get("anthropic_api_key"):
                result = self._apply_llm(result)
            st.session_state["ocr_result"] = result
            st.session_state["edited_text"] = result["text"]
            placeholder.empty()
            st.rerun()
        except Exception as e:
            placeholder.empty()
            st.error(f"🚨 Agent error: {e}")
            import traceback
            with st.expander("Stack trace"):
                st.code(traceback.format_exc())

    def _apply_llm(self, result: dict) -> dict:
        try:
            import anthropic
            key = st.session_state.get("anthropic_api_key", "")
            if not key:
                return result
            client = anthropic.Anthropic(api_key=key)
            msg = client.messages.create(
                model="claude-sonnet-4-20250514", max_tokens=1000,
                messages=[{"role": "user", "content":
                           f"Fix OCR artefacts. Correct misread characters. Preserve meaning. Return corrected text only.\n\n{result['text'][:3000]}"}])
            result["text"] = msg.content[0].text
            result["decision_log"].append("[LLM] Claude correction applied.")
        except Exception as e:
            result["decision_log"].append(f"[LLM] Skipped: {e}")
        return result

    def _make_docx(self, text: str, r: dict) -> bytes:
        from docx import Document
        from docx.shared import Pt, RGBColor
        doc = Document()
        h = doc.add_heading("OctoRead — Extracted Text", level=1)
        h.runs[0].font.color.rgb = RGBColor(0xE8, 0x62, 0x3C)
        meta = doc.add_paragraph()
        meta.add_run(f"Strategy: {r.get('strategy_used','—')} | Confidence: {r.get('confidence',0):.1f}% | Words: {r.get('word_count',0)}").font.size = Pt(9)
        doc.add_paragraph()
        for line in text.split("\n"):
            doc.add_paragraph(line)
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
